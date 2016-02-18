'''this file searches the journal publications PDF files for the acknowledgement
to the CSA'''
##################
from sqlalchemy import create_engine, asc
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from acepy.getquarterlystats import pubrecord
from datetime import datetime as dt
from docx import Document

'''htmlreplacements = {'&eacute;':'',
                    '&Eacute;':'',
                    '&ouml;':'',
                    '&Ouml;':''}

def removehtml(pub):
    for attribute in ['lead_author_surname','lead_author_given_names','title','coauthors']:
        attstr = pub.__dict__[attribute]
'''

#port = '5432' #for production
port = '5433' #for testing
webdataEngine = create_engine('postgresql://rchughes:h2olo2h2o@localhost:'+port+'/web_data')
webdataBase = declarative_base(webdataEngine)
startdate = dt(2015,1,1,0,0,0)
enddate = dt(2015,12,31,23,59,59)

metadata = webdataBase.metadata
webdataSession = sessionmaker(bind=webdataEngine)
mywebdatasession = webdataSession()

pubs2015 = mywebdatasession.query(pubrecord).\
    filter(pubrecord.pub_date>startdate).\
    filter(pubrecord.pub_date<enddate).\
    order_by(asc(pubrecord.lead_author_surname)).all()

manualPubsWithCSAref = []
with open('searchOutput_2016-02-12.txt','r') as f:
    for line in f:
        manualPubsWithCSAref.append(line.strip())

pubsWithCSAref = []
otherpubs = []

for pub in pubs2015:
    if pub.filename in manualPubsWithCSAref:
        pubsWithCSAref.append(pub)
    else:
        otherpubs.append(pub)

#outputting publications to a word document
document = Document()
document.add_heading('Performance Indicators for '+str(startdate.year),0)

document.add_heading('PEER Reviewed Publications',1)
document.add_heading('Acknowledges Canadian Space Agency',2)
if len(pubsWithCSAref)>0:
    for i,pub in enumerate(pubsWithCSAref):
        if i==0: pub.lead_author_surname = 'Flinstone'
        p = document.add_paragraph(pub.lead_author_surname+', '+\
            pub.lead_author_given_names+', '+\
            pub.coauthors+'. ', style='List')
        p.add_run(pub.title).underline=True
        p.add_run('. ')
        p.add_run(pub.journal_name+' ').italic=True
        p.add_run(pub.volume).bold=True
        p.add_run(', '+pub.pages+'.\n')
else:
    document.add_paragraph('N/A',style='Normal')

document.add_heading('No Canadian Space Agency Acknowledgement',2)
if len(otherpubs)>0:
    for pub in otherpubs:
        p = document.add_paragraph(pub.lead_author_surname+', '+\
            pub.lead_author_given_names+', '+\
            pub.coauthors+'. ', style='List')
        p.add_run(pub.title).underline=True
        p.add_run('. ')
        p.add_run(pub.journal_name+' ').italic=True
        p.add_run(pub.volume).bold=True
        p.add_run(', '+pub.pages+'.\n')
else:
    document.add_paragraph('N/A',style='Normal')


document.add_heading('Presentations',1)
#document.add_paragraph('sample item', style='ListNumber')

document.save('PerformanceIndicatorPublications_'+str(startdate.year)+'.docx')

'''import glob
import re
from pdf2txtfunc import searchpdf

def searchyear(year):
    #pubfolder = '/tankhome/rchughes/acebox2/publications/'
    pubfolder = './'
    pubfiles = glob.glob(pubfolder+str(year)+'/*.pdf') #find all publications
    termsfound = {}

    for file in pubfiles:
        #folderfile = file.split('/')[-2]+'/'+file.split('/')[-1]
        #print '*** Converting PDF to text: '+folderfile+'...'
        #pdftext = pdf2txt(file)
        #print '*** Searching file: '+folderfile+'...'
        #files_with_reference = []
        #for searchterm in ['Space', 'Agency','CSA','anadian']:
        #    myregex = re.escape(searchterm)
        #    m = re.search(r'('+myregex+r')',pdftext)
        #    if m is not None:
        #        files_with_reference.append(file)
        #        print 'File '+folderfile+' contains term \''+searchterm+'\'.'

        searchterms = ['Space','Agency','CSA','anadian']
        termsfound[file.split('/')[-1]] = searchpdf(file,searchterms)

    return termsfound

if __name__ == "__main__":
    searchyear(2015)'''
